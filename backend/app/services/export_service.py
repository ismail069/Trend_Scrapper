from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models import ExportRequest

ACCENT = HexColor("#2ca83a")


def build_pdf(payload: ExportRequest) -> BytesIO:
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AppTitle",
        parent=styles["Title"],
        textColor=ACCENT,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    metadata_style = ParagraphStyle(
        "Metadata",
        parent=styles["BodyText"],
        fontSize=8,
        leading=10,
        textColor=HexColor("#657067"),
    )
    story = [
        Paragraph("Feed Tren Scrapper", title_style),
        Paragraph(f"<b>Search:</b> {escape(payload.prompt)}", styles["BodyText"]),
        Paragraph(f"<b>Category:</b> {escape(payload.category)}", styles["BodyText"]),
        Paragraph(
            f"<b>Datetime:</b> {payload.searched_at.strftime('%Y-%m-%d %H:%M %Z')}",
            styles["BodyText"],
        ),
        Spacer(1, 14),
    ]
    if not payload.results:
        story.append(Paragraph("No results were available.", styles["Italic"]))
    for index, result in enumerate(payload.results, start=1):
        story.extend(
            [
                Paragraph(f"{index}. {escape(result.title)}", styles["Heading3"]),
                Paragraph(escape(result.summary), styles["BodyText"]),
                Paragraph(
                    escape(
                        " | ".join(
                            value
                            for value in [
                                result.source_name,
                                str(result.source_url) if result.source_url else None,
                                (
                                    result.published_at.strftime("%Y-%m-%d")
                                    if result.published_at
                                    else None
                                ),
                            ]
                            if value
                        )
                    ),
                    metadata_style,
                ),
                Spacer(1, 10),
            ]
        )
    document.build(story)
    buffer.seek(0)
    return buffer


def build_docx(payload: ExportRequest) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    title = document.add_heading("Feed Tren Scrapper", level=0)
    title.style = document.styles["Title"]
    document.add_paragraph(f"Search: {payload.prompt}")
    document.add_paragraph(f"Category: {payload.category}")
    document.add_paragraph(
        f"Datetime: {payload.searched_at.strftime('%Y-%m-%d %H:%M %Z')}"
    )
    if not payload.results:
        document.add_paragraph("No results were available.")
    for index, result in enumerate(payload.results, start=1):
        document.add_heading(f"{index}. {result.title}", level=2)
        document.add_paragraph(result.summary)
        metadata = " | ".join(
            value
            for value in [
                result.source_name,
                str(result.source_url) if result.source_url else None,
                result.published_at.strftime("%Y-%m-%d") if result.published_at else None,
            ]
            if value
        )
        if metadata:
            document.add_paragraph(metadata)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
